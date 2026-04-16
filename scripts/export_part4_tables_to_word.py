from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "scripts" / "part4_comparative.knit.md"
OUTPUT_DOCX = ROOT / "docs" / "part4_results_report.docx"
LEGACY_OUTPUT_DOCX = ROOT / "docs" / "part4_result_tables.docx"
FIGURES_DIR = ROOT / "figures"

TABLE_SPECS = [
    {
        "number": 1,
        "caption": "Coverage check for Part 4 analysis inputs.",
        "analysis": (
            "Table 1 demonstrates that the retained Part 4 dataset achieves a high level of analytic completeness, "
            "which strengthens the comparability of inter-university travel patterns. The remaining unmatched records "
            "should be interpreted as a limited but non-negligible source of measurement uncertainty, particularly for "
            "low-frequency destinations and edge cases in station-topic linkage."
        ),
        "section": "raw",
    },
    {
        "number": 2,
        "caption": "Campus station counts and total trips in the study period.",
        "analysis": (
            "Table 2 shows that the four campuses enter the analysis with markedly different transit footprints and trip "
            "volumes. These structural asymmetries are analytically important because they shape both the breadth of the "
            "reachable destination set and the observed dependence on bus versus rail access."
        ),
        "section": "raw",
    },
    {
        "number": 8,
        "caption": "Weighted Jaccard similarity based on topic-share distributions.",
        "analysis": (
            "Table 8 quantifies the degree of overlap in topic composition between university communities. Higher weighted "
            "Jaccard values indicate convergence in destination-function structure, whereas lower values signal more distinct "
            "urban consumption and mobility profiles."
        ),
        "section": "topic",
    },
    {
        "number": 13,
        "caption": "Universities with significant positive deviations in topic preference.",
        "analysis": (
            "Table 13 identifies the universities that significantly over-index in specific destination functions relative to "
            "the pooled expectation. These positive deviations clarify which functional orientations are not merely descriptive "
            "patterns, but statistically salient expressions of differentiated university-related mobility."
        ),
        "section": "topic",
    },
]

FIGURE_SPECS = [
    {
        "number": 1,
        "title": "Travel distance distributions by mode for NUS, NTU, SMU, and SUTD.",
        "path": FIGURES_DIR / "part4_fig2_distance_boxplot.png",
        "analysis": (
            "Figure 1 reveals substantial variation in trip-distance structure across both universities and modes. Bus trips "
            "remain comparatively localised, whereas MRT trips extend the effective spatial field of activity, indicating a "
            "clear scaling up of destination reach once communities gain direct access to the rail network."
        ),
        "section": "raw",
        "width": 6.6,
    },
    {
        "number": 2,
        "title": "Destinations ranked by trip frequency for each university community.",
        "path": FIGURES_DIR / "part4_fig3_top20_bus_destinations.png",
        "analysis": (
            "Figure 2 foregrounds the strong concentration of bus-based mobility into a limited hierarchy of recurrent destinations. "
            "At the same time, the destination rankings vary meaningfully across campuses, suggesting that bus travel retains a more "
            "place-specific and institutionally embedded geography than aggregate citywide flow metrics would imply."
        ),
        "section": "raw",
        "width": 6.6,
    },
    {
        "number": 3,
        "title": "Ranked by trip frequency across university communities (NTU has no MRT-origin trips).",
        "path": FIGURES_DIR / "part4_fig4_top20_mrt_destinations.png",
        "analysis": (
            "Figure 3 shows that MRT-linked destination systems are both more regionally extensive and more sharply ordered than "
            "their bus counterparts. The absence of NTU MRT-origin trips is itself analytically telling, underscoring the infrastructural "
            "asymmetry that differentiates the mobility options available to the four university communities."
        ),
        "section": "raw",
        "width": 6.6,
    },
    {
        "number": 4,
        "title": "Spatial reach of trips across distance bands by university.",
        "path": FIGURES_DIR / "part4_fig6_distance_rings.png",
        "analysis": (
            "Figure 4 translates travel distance into an interpretable territorial profile by showing how each university community's trips "
            "are distributed across nested distance bands. The resulting pattern clarifies which campuses are dominated by proximate urban "
            "activity and which display a broader metropolitan catchment."
        ),
        "section": "raw",
        "width": 6.4,
    },
    {
        "number": 5,
        "title": "Compares destination intensity patterns of NUS, NTU, SMU, and SUTD across Singapore.",
        "path": FIGURES_DIR / "part5_destination_density_map.png",
        "analysis": (
            "Figure 5 captures the broader spatial morphology of destination intensity by smoothing point-based trip concentrations across "
            "Singapore. The contour structure demonstrates that university-linked mobility is not simply dispersed across the city, but forms "
            "distinct clusters whose territorial emphasis differs markedly among campuses."
        ),
        "section": "raw",
        "width": 6.6,
    },
    {
        "number": 6,
        "title": "Distribution of trips from universities to destinations across Singapore.",
        "path": FIGURES_DIR / "part5_destination_flow_map.png",
        "analysis": (
            "Figure 6 complements the density surface by restoring the directional logic of the origin-destination system. The map makes visible "
            "how each university anchors a distinctive field of outbound connections, with overlapping yet clearly differentiated corridors of urban engagement."
        ),
        "section": "raw",
        "width": 6.6,
    },
    {
        "number": 7,
        "title": "Distribution of trips to urban functions across university communities.",
        "path": FIGURES_DIR / "part4_fig7_topic_share_bar.png",
        "analysis": (
            "Figure 7 compares the proportional composition of trips across urban functions and therefore shifts the analysis from where students travel "
            "to what kinds of places they preferentially access. The stacked profile highlights the balance between shared metropolitan routines and distinctive "
            "institution-specific destination preferences."
        ),
        "section": "topic",
        "width": 6.1,
    },
    {
        "number": 8,
        "title": "Distribution of trips across regions and urban functions by university.",
        "path": FIGURES_DIR / "part4_fig5_sankey_university_region_topic.png",
        "analysis": (
            "Figure 8 integrates region and urban function within a single relational diagram, revealing how destination structure is jointly organised by "
            "territorial location and functional character. This multi-stage representation is especially useful for identifying whether university communities "
            "share the same functions in different regions or, conversely, converge spatially while diverging functionally."
        ),
        "section": "topic",
        "width": 6.6,
    },
]


def set_font(run, name="Times New Roman", size=10.5, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, space_before=0, line_spacing=1.15):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_do_not_compress_pictures(document):
    settings = document.settings.element
    existing = settings.find(qn("w:doNotCompressPictures"))
    if existing is None:
        settings.append(OxmlElement("w:doNotCompressPictures"))


def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tc_borders = tcPr.first_child_found_in("w:tcBorders")
    if tc_borders is not None:
        tcPr.remove(tc_borders)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tc_borders = tcPr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def parse_pipe_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            break
        rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
    if len(rows) < 2:
        return None
    return rows[0], rows[2:]


def extract_tables(markdown_text):
    lines = markdown_text.splitlines()
    tables = []
    i = 0
    while i < len(lines):
        match = re.match(r"Table:\s+(.*)", lines[i].strip())
        if not match:
            i += 1
            continue

        caption = match.group(1).strip()
        j = i + 1
        while j < len(lines) and not lines[j].strip().startswith("|"):
            j += 1
        if j >= len(lines):
            break

        table_lines = []
        while j < len(lines) and lines[j].strip().startswith("|"):
            table_lines.append(lines[j])
            j += 1

        parsed = parse_pipe_table(table_lines)
        if parsed is not None:
            header, body = parsed
            tables.append({"caption": caption, "header": header, "body": body})
        i = j
    return tables


def add_section_heading(document, text):
    para = document.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
    run = para.add_run(text)
    set_font(run, size=12, bold=True)


def add_caption(document, label, number, title, is_table=False):
    para = document.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    lead = para.add_run(f"{label} {number}. ")
    set_font(lead, size=11, bold=True)
    rest = para.add_run(title)
    set_font(rest, size=11, bold=False)
    if is_table:
        para.paragraph_format.keep_with_next = True
    return para


def add_analysis_paragraph(document, text):
    para = document.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=2, space_after=10, line_spacing=1.2)
    run = para.add_run(text)
    set_font(run, size=10.5)


def add_three_line_table(document, number, caption, header, body):
    add_caption(document, "Table", number, caption, is_table=True)

    table = document.add_table(rows=1 + len(body), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for col_idx, value in enumerate(header):
        cell = table.cell(0, col_idx)
        cell.text = value
        para = cell.paragraphs[0]
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in para.runs:
            set_font(run, size=10.5, bold=True)

    for row_idx, row in enumerate(body, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            para = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(para, alignment=align)
            for run in para.runs:
                set_font(run, size=10.5)

    border_style = {"val": "single", "sz": 8, "space": 0, "color": "000000"}
    no_border = {"val": "nil"}

    n_rows = 1 + len(body)
    n_cols = len(header)

    for r in range(n_rows):
        for c in range(n_cols):
            cell = table.cell(r, c)
            clear_cell_borders(cell)
            set_cell_border(
                cell,
                top=no_border,
                bottom=no_border,
                left=no_border,
                right=no_border,
                insideH=no_border,
                insideV=no_border,
            )

    for c in range(n_cols):
        set_cell_border(table.cell(0, c), top=border_style, bottom=border_style)
        set_cell_border(table.cell(n_rows - 1, c), bottom=border_style)


def add_figure(document, number, title, image_path, width_inches):
    if not image_path.exists():
        raise FileNotFoundError(f"Missing figure: {image_path}")

    picture_para = document.add_paragraph()
    set_paragraph_format(picture_para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=3)
    run = picture_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    add_caption(document, "Figure", number, title)


def build_document(tables_by_caption):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    add_do_not_compress_pictures(doc)

    title = doc.add_paragraph()
    set_paragraph_format(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    run = title.add_run("Part 4 Results Report")
    set_font(run, size=14, bold=True)

    intro = doc.add_paragraph()
    set_paragraph_format(intro, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.2)
    intro_run = intro.add_run(
        "This report reorganises the Part 4 outputs into a publication-oriented sequence, moving from raw destination structure "
        "to destination-topic interpretation. Tables are presented in three-line format, and figures are embedded at full resolution "
        "for direct use in an academic manuscript."
    )
    set_font(intro_run, size=10.5)

    add_section_heading(doc, "Raw Destination Patterns")
    for spec in [item for item in TABLE_SPECS if item["section"] == "raw"]:
        table = tables_by_caption.get(spec["caption"])
        if table is None:
            raise KeyError(f"Required table not found: {spec['caption']}")
        add_three_line_table(doc, spec["number"], spec["caption"], table["header"], table["body"])
        add_analysis_paragraph(doc, spec["analysis"])

    for spec in [item for item in FIGURE_SPECS if item["section"] == "raw"]:
        add_figure(doc, spec["number"], spec["title"], spec["path"], spec["width"])
        add_analysis_paragraph(doc, spec["analysis"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_section_heading(doc, "Destination Topic Patterns")
    for spec in [item for item in TABLE_SPECS if item["section"] == "topic"]:
        if spec["number"] == 8:
            table = tables_by_caption.get(spec["caption"])
            if table is None:
                raise KeyError(f"Required table not found: {spec['caption']}")
            add_three_line_table(doc, spec["number"], spec["caption"], table["header"], table["body"])
            add_analysis_paragraph(doc, spec["analysis"])

            for fig_spec in [item for item in FIGURE_SPECS if item["section"] == "topic"]:
                add_figure(doc, fig_spec["number"], fig_spec["title"], fig_spec["path"], fig_spec["width"])
                add_analysis_paragraph(doc, fig_spec["analysis"])

        elif spec["number"] == 13:
            table = tables_by_caption.get(spec["caption"])
            if table is None:
                raise KeyError(f"Required table not found: {spec['caption']}")
            add_three_line_table(doc, spec["number"], spec["caption"], table["header"], table["body"])
            add_analysis_paragraph(doc, spec["analysis"])

    doc.save(OUTPUT_DOCX)
    doc.save(LEGACY_OUTPUT_DOCX)


def main():
    markdown_text = SOURCE_MD.read_text(encoding="utf-8")
    tables = extract_tables(markdown_text)
    tables_by_caption = {item["caption"]: item for item in tables}
    build_document(tables_by_caption)
    print(f"Saved report to {OUTPUT_DOCX}")
    print(f"Saved report to {LEGACY_OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
